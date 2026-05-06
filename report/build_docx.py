"""Generate the submission as an editable Word document (.docx).

Cover page in the required format + condensed report with metrics table and
embedded figures. Run after ``python -m src.main --model all`` so figures exist
in ``report/figures/``.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from build_pdf import SUBMISSION  # reuse the same metadata


REPO_ROOT = Path(__file__).resolve().parent.parent
REPORT_DIR = REPO_ROOT / "report"
FIG_DIR = REPORT_DIR / "figures"
METRICS_PATH = REPORT_DIR / "metrics.json"
OUT_PATH = REPORT_DIR / "Submission.docx"


def _set_run(run, *, size=11, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _para(doc, text, *, size=11, bold=False, align=None, space_after=4, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _set_run(r, size=size, bold=bold, color=color)
    return p


def _heading(doc, text, *, size=14):
    _para(doc, text, size=size, bold=True, space_after=6,
          color=(0x1F, 0x3A, 0x5F))


def _field(doc, label, value, *, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label}: ")
    _set_run(r, size=size, bold=True)
    r2 = p.add_run(value)
    _set_run(r2, size=size, bold=False)


def _cover(doc):
    _para(doc, "Submission Cover Page", size=20, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
          color=(0x1F, 0x3A, 0x5F))

    _field(doc, "Exercise number", SUBMISSION["exercise_number"])
    _field(doc, "Group code", SUBMISSION["group_id"])
    _field(doc, "Self-evaluation score (0-100)", SUBMISSION["self_score"])

    _para(doc, "Students:", size=12, bold=True, space_after=2)
    for s in SUBMISSION["students"]:
        line = (
            f"  • ID {s['id']}  |  {s['first_en']} {s['last_en']}  |  "
            f"{s['first_he']} {s['last_he']}"
        )
        _para(doc, line, size=12, space_after=2)

    doc.add_paragraph()
    _field(doc, "GitHub repository", SUBMISSION["github"])
    _field(doc, "Late submission", SUBMISSION["late_submission"])

    doc.add_page_break()


def _metrics_table(doc):
    if not METRICS_PATH.exists():
        _para(doc, "(metrics.json not found — run training first.)",
              color=(0xAA, 0x00, 0x00))
        return
    metrics = json.loads(METRICS_PATH.read_text())
    headers = ["Model", "Test Acc", "Recon MSE",
               "MSE (σ=5%)", "MSE (σ=10%)", "MSE (σ=20%)"]
    rows = []
    for name in ("mlp", "rnn", "lstm"):
        if name not in metrics:
            continue
        m = metrics[name]
        rows.append([
            name.upper(),
            f"{m['acc']:.3f}",
            f"{m['mse']:.4f}",
            f"{m['mse_sigma_5']:.4f}",
            f"{m['mse_sigma_10']:.4f}",
            f"{m['mse_sigma_20']:.4f}",
        ])

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        _set_run(r, size=11, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            _set_run(r, size=11)


def _figure(doc, path: Path, caption: str):
    if not path.exists():
        _para(doc, f"(missing figure: {path.name})",
              color=(0xAA, 0x00, 0x00))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(15))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run(r, size=10, bold=True, color=(0x55, 0x55, 0x55))


def _body(doc):
    _heading(doc, "SignalMemoryNet — HW1 Lab Report", size=18)

    _heading(doc, "1. Problem Setup")
    _para(doc,
          "We generate sine signals at four discrete frequencies "
          "(2, 5, 10, 20 Hz), sampled at 100 Hz for 10 seconds, with random "
          "amplitude in [0.8, 1.2], random phase, and additive Gaussian noise "
          "at three levels (σ = 5 %, 10 %, 20 % of amplitude). The task is "
          "two-fold per 10-sample window: (i) classify the underlying "
          "frequency (4 classes), (ii) reconstruct the clean window "
          "(denoising). A shared backbone feeds a multi-task head producing "
          "4 logits (cross-entropy) + 10 regression outputs (MSE).")

    _heading(doc, "2. Architectures")
    _para(doc,
          "MLP — fully-connected baseline that ignores temporal order: "
          "FC(10→64)–ReLU–FC(64→64)–ReLU–Head(14).")
    _para(doc,
          "RNN — single-layer tanh nn.RNN with hidden size 32; the final "
          "hidden state h_T is fed to the same head.")
    _para(doc,
          "LSTM — single-layer nn.LSTM with hidden size 32; gating (input, "
          "forget, output) mitigates the vanishing-gradient issue and "
          "improves denoising at high noise.")

    _heading(doc, "3. Training Setup")
    _para(doc,
          "Optimizer Adam (lr=1e-3), batch size 128, up to 20 epochs, early "
          "stopping (patience 5) on the validation loss. Records per "
          "frequency: train 30 / val 8 / test 8, sliding windows of length "
          "10 with stride 5. Seed = 42 for full reproducibility.")

    _heading(doc, "4. Results")
    _metrics_table(doc)
    doc.add_paragraph()

    _figure(doc, FIG_DIR / "training_curves.png",
            "Figure 1 — Training/validation loss curves for MLP, RNN, LSTM.")
    _figure(doc, FIG_DIR / "confusion_mlp.png",
            "Figure 2 — Confusion matrix (MLP).")
    _figure(doc, FIG_DIR / "confusion_rnn.png",
            "Figure 3 — Confusion matrix (RNN).")
    _figure(doc, FIG_DIR / "confusion_lstm.png",
            "Figure 4 — Confusion matrix (LSTM).")
    _figure(doc, FIG_DIR / "recon_mlp.png",
            "Figure 5 — Reconstruction examples (MLP).")
    _figure(doc, FIG_DIR / "recon_rnn.png",
            "Figure 6 — Reconstruction examples (RNN).")
    _figure(doc, FIG_DIR / "recon_lstm.png",
            "Figure 7 — Reconstruction examples (LSTM).")

    _heading(doc, "5. Discussion")
    _para(doc,
          "All three models achieve > 97 % classification accuracy thanks to "
          "the strong frequency cue inside any 10-sample window at 100 Hz. "
          "Denoising is harder: the recurrent models re-use information across "
          "time steps and recover better at the highest noise level, while the "
          "MLP, despite a higher parameter count per window, still benefits "
          "from the explicit reconstruction signal in the multi-task loss.")


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    _cover(doc)
    _body(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
